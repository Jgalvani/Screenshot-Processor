"""Word document reader utility for extracting URLs."""

import re
from pathlib import Path
from docx import Document


class WordReader:
    """Reads and extracts URLs from Word documents."""

    URL_PATTERN = re.compile(
        r"https?://(?:[-\w.]|(?:%[\da-fA-F]{2}))+"
        r"(?:/[-\w%!$&'()*+,./:;=?@~#]*)*"
    )

    def __init__(self, file_path: str | Path):
        """
        Initialize WordReader with a document path.

        Args:
            file_path: Path to the Word document (.docx)
        """
        self.file_path = Path(file_path)
        self._validate_file()

    def _validate_file(self) -> None:
        """Validate that the file exists and is a .docx file."""
        if not self.file_path.exists():
            raise FileNotFoundError(f"File not found: {self.file_path}")
        if self.file_path.suffix.lower() != ".docx":
            raise ValueError(f"Expected .docx file, got: {self.file_path.suffix}")

    def extract_urls(self) -> list[str]:
        """
        Extract all URLs from the Word document.

        Returns:
            List of unique URLs found in the document
        """
        document = Document(self.file_path)
        urls = set()

        # Extract from paragraphs
        for paragraph in document.paragraphs:
            found_urls = self.URL_PATTERN.findall(paragraph.text)
            urls.update(found_urls)

        # Extract from tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    found_urls = self.URL_PATTERN.findall(cell.text)
                    urls.update(found_urls) 

        return list(urls)

    def _extract_hyperlinks_from_xml(self, xml_content: str) -> list[str]:
        """Extract hyperlink URLs from XML content."""
        return self.URL_PATTERN.findall(xml_content)

    def get_urls_with_context(self) -> list[dict]:
        """
        Extract URLs with surrounding text context.

        Returns:
            List of dicts with 'url' and 'context' keys
        """
        document = Document(self.file_path)
        results = []

        for paragraph in document.paragraphs:
            found_urls = self.URL_PATTERN.findall(paragraph.text)
            for url in found_urls:
                results.append({
                    "url": url,
                    "context": paragraph.text.strip()
                })

        return results
